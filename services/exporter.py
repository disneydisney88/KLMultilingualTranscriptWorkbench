from __future__ import annotations

import json
from pathlib import Path
from typing import Sequence

from docx import Document

from services.common import (
    TranscriptResult,
    TranscriptSegment,
    ensure_dir,
    seconds_to_timestamp,
    seconds_to_vtt_timestamp,
)


DEFAULT_EXPORT_FORMATS = ("txt", "docx")


def _normalize_export_formats(formats: Sequence[str] | None) -> list[str]:
    if not formats:
        return list(DEFAULT_EXPORT_FORMATS)
    ordered: list[str] = []
    seen: set[str] = set()
    for item in formats:
        value = str(item).strip().lower()
        if value and value not in seen:
            ordered.append(value)
            seen.add(value)
    return ordered or list(DEFAULT_EXPORT_FORMATS)


def _segment_prefix(segment: TranscriptSegment) -> str:
    speaker = f"{segment.speaker}: " if segment.speaker else ""
    return f"[{seconds_to_timestamp(segment.start, milliseconds=False)} - {seconds_to_timestamp(segment.end, milliseconds=False)}] {speaker}"


def _build_transcript_text(result: TranscriptResult) -> str:
    lines = [
        f"Job ID: {result.job_id}",
        f"Source: {result.source_type}",
        f"Title: {result.media_title or ''}",
        f"Original File: {result.original_filename or ''}",
        f"Language: {result.language}",
        f"Model: {result.model_size}",
        "",
    ]
    for segment in result.segments:
        text = segment.text_corrected.strip()
        if not text:
            continue
        lines.append(f"{_segment_prefix(segment)}{text}")
    return "\n".join(lines).strip()


def _build_srt(result: TranscriptResult) -> str:
    blocks: list[str] = []
    for index, segment in enumerate(result.segments, start=1):
        speaker_prefix = f"{segment.speaker}: " if segment.speaker else ""
        blocks.append(
            "\n".join(
                [
                    str(index),
                    f"{seconds_to_timestamp(segment.start, milliseconds=True)} --> {seconds_to_timestamp(segment.end, milliseconds=True)}",
                    f"{speaker_prefix}{segment.text_corrected}".strip(),
                ]
            )
        )
    return "\n\n".join(blocks).strip()


def _build_vtt(result: TranscriptResult) -> str:
    blocks: list[str] = ["WEBVTT", ""]
    for segment in result.segments:
        speaker_prefix = f"{segment.speaker}: " if segment.speaker else ""
        blocks.append(
            "\n".join(
                [
                    f"{seconds_to_vtt_timestamp(segment.start)} --> {seconds_to_vtt_timestamp(segment.end)}",
                    f"{speaker_prefix}{segment.text_corrected}".strip(),
                    "",
                ]
            )
        )
    return "\n".join(blocks).strip() + "\n"


def _write_text(path: Path, content: str) -> Path:
    path.write_text(content.strip() + "\n", encoding="utf-8")
    return path


def _write_docx_transcript(path: Path, result: TranscriptResult) -> Path:
    doc = Document()
    doc.add_heading("撱?閰梢?蝔踹極雿", level=0)
    meta = doc.add_table(rows=0, cols=2)
    meta.style = "Table Grid"

    def add_meta(label: str, value: object) -> None:
        row = meta.add_row().cells
        row[0].text = str(label)
        row[1].text = str(value)

    add_meta("Job ID", result.job_id)
    add_meta("Source", result.source_type)
    add_meta("Title", result.media_title or "")
    add_meta("Original File", result.original_filename or "")
    add_meta("Language", result.language)
    add_meta("Model", result.model_size)
    if result.duration_seconds is not None:
        add_meta("Duration (sec)", f"{result.duration_seconds:.1f}")

    doc.add_paragraph("")
    doc.add_heading("??蝔?, level=1)
    for segment in result.segments:
        if not segment.text_corrected.strip():
            continue
        prefix = _segment_prefix(segment)
        doc.add_paragraph(f"{prefix}{segment.text_corrected}")

    doc.save(path)
    return path


def _write_summary_docx(path: Path, result: TranscriptResult) -> Path:
    doc = Document()
    doc.add_heading("Summary Notes", level=0)
    doc.add_paragraph(result.summary or "No summary available.")

    doc.add_heading("Key Moments", level=1)
    for segment in result.segments[:20]:
        text = segment.text_corrected.strip()
        if not text:
            continue
        speaker = f"{segment.speaker}: " if segment.speaker else ""
        doc.add_paragraph(
            f"{seconds_to_timestamp(segment.start, milliseconds=False)} - {seconds_to_timestamp(segment.end, milliseconds=False)} {speaker}{text}",
            style="List Bullet",
        )

    if result.glossary_changes:
        doc.add_heading("Glossary Applied", level=1)
        for change in result.glossary_changes:
            doc.add_paragraph(f"{change.pattern} -> {change.replacement} ({change.count})", style="List Bullet")

    doc.save(path)
    return path


def export_transcript_bundle(
    result: TranscriptResult,
    output_dir: Path,
    formats: Sequence[str] | None = None,
) -> list[Path]:
    ensure_dir(output_dir)
    selected = _normalize_export_formats(formats)
    files: list[Path] = []

    if "txt" in selected:
        files.append(_write_text(output_dir / "transcript.txt", _build_transcript_text(result)))
    if "srt" in selected:
        files.append(_write_text(output_dir / "transcript.srt", _build_srt(result)))
    if "vtt" in selected:
        files.append(_write_text(output_dir / "transcript.vtt", _build_vtt(result)))
    if "json" in selected:
        files.append(_write_text(output_dir / "transcript.json", json.dumps(result.to_dict(), ensure_ascii=False, indent=2)))
    if "docx" in selected:
        files.append(_write_docx_transcript(output_dir / "transcript.docx", result))
    if "summary_docx" in selected:
        files.append(_write_summary_docx(output_dir / "summary_notes.docx", result))

    return files
